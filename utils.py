import os
import sqlite3
import json
from datetime import datetime
from docx import Document
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from litellm import completion

# Carrega variáveis de ambiente do arquivo .env
load_dotenv() 
pasta_relatorios = os.getenv("PASTA_DOWNLOADS", os.path.expanduser("~\\Downloads"))
modelo = os.getenv('modelo_selecionado')
api_key = os.getenv('OPENAI_API_KEY')
prompt = os.getenv('prompt')

# Funções para a leitura dos arquivos PDF, DOCX e TXT
def ler_pdf(caminho_pdf):
    reader = PdfReader(caminho_pdf)
    texto = ''
    for page in reader.pages:
        texto += page.extract_text() + '\n'
    return texto

def ler_docx(caminho_docx):
    doc = Document(caminho_docx)
    texto = ''
    for paragrafo in doc.paragraphs:
        texto += paragrafo.text + '\n'
    return texto

def ler_txt(caminho_txt):
    with open(caminho_txt, 'r', encoding='utf-8') as file:
        texto = file.read()
    return texto

# Função para limpar o texto extraído
import re

def limpar_texto(texto):
    if not texto:
        return ""
    # 1. Remove espaços no início e fim do documento
    texto = texto.strip()
    # 2. Transforma múltiplos espaços ou tabulações em um único espaço
    texto = re.sub(r'[ \t]+', ' ', texto)
    # 3. Substitui quebras de linha múltiplas por uma única quebra (\n)
    texto = re.sub(r'\n\s*\n+', '\n', texto)
    # 4. Limpa espaços remanescentes no início e fim de cada linha individual
    linhas = [linha.strip() for linha in texto.split('\n') if linha.strip()]
    return '\n'.join(linhas)

# Função para obter o texto do arquivo com base na extensão
def obter_texto(caminho_arquivo):
    extensao = os.path.splitext(caminho_arquivo)[1].lower()
    texto = '' 
    if extensao == '.pdf':
        texto = ler_pdf(caminho_arquivo)
    elif extensao == '.docx':
        texto = ler_docx(caminho_arquivo)
    elif extensao == '.txt':
        texto = ler_txt(caminho_arquivo)
    else:
        texto = ''
    texto = limpar_texto(texto)
    return texto

def analisar_conteudo(texto, prompt, modelo, api_key):
    print(f"Usando o modelo: {modelo}")
    try:
        instrucao_formato = (
            "Responda exclusivamente em formato JSON com as seguintes chaves: "
            "'assunto' (string correspondendo a uma única frase) e 'resumo' (string contendo os 4 parágrafos solicitados)."
        )

        resultado = completion(
            model=modelo,
            messages=[
                {"role": "system", "content": "Você é um analista jurídico."},
                {"role": "system", "content": f"{prompt}\n\n{instrucao_formato}"},
                {"role": "user", "content": (
                    "Analise o texto abaixo e retorne o resumo em exatamente 4 parágrafos:\n"
                    "1. Visão geral e estrutura do documento.\n"
                    "2. Questões tratadas em cada tópico.\n"
                    "3. Conclusão bem fundamentada sobre o teor.\n"
                    "4. Pendências explicitamente mencionadas.\n\n"
                    f"**Texto fornecido:**\n{texto}"
                )},
            ],
            api_key=api_key,
            response_format={"type": "json_object"} 
        )

        conteudo_raw = resultado.choices[0].message.content
        dados = json.loads(conteudo_raw)
        
        # Extrai os dados do JSON retornado pela IA
        assunto = dados.get('assunto', 'Assunto não identificado')
        resumo = dados.get('resumo', 'Resumo não gerado')

        return (assunto, resumo) 

    except Exception as e:
        print(f"Erro na análise: {e}")
        return ("Erro", "Não foi possível processar o conteúdo.")

def gerar_titulo_e_resumo(caminho_arquivo):
    texto = obter_texto(caminho_arquivo)
    if not texto:
        return "Erro", "Não foi possível ler o arquivo."
    assunto, resumo = analisar_conteudo(texto, prompt, modelo, api_key)
    return assunto, resumo

def exportar_relatorio(lista_dados):
    try:
        doc = Document()
        doc.add_heading('PJ Docs', 0)

        # 1. Agrupamento por procedimento
        agrupados = {}
        for item in lista_dados:
            proc = item['procedimento']
            if proc not in agrupados:
                agrupados[proc] = []
            agrupados[proc].append(item)

        # 2. Construção da hierarquia no documento
        for procedimento, documentos in agrupados.items():
            doc.add_heading(f"Autos/expediente: {procedimento}", level=1)
            doc.add_paragraph("")  
            
            for doc_info in documentos:
                # --- DATA DO CADASTRO ---
                p_data = doc.add_paragraph()
                p_data.add_run('Data: ').bold = True
                
                data_original = doc_info['data']
                try:
                    data_formatada = datetime.strptime(data_original, "%Y-%m-%d").strftime("%d/%m/%Y")
                except Exception:
                    data_formatada = data_original

                p_data.add_run(data_formatada)

                # --- TÍTULO DO DOCUMENTO ---
                p_titulo = doc.add_paragraph()
                p_titulo.add_run('Documento: ').bold = True
                p_titulo.add_run(doc_info['titulo'])
                
                # --- RESUMO ---
                p_resumo = doc.add_paragraph()
                p_resumo.add_run('Teor: ').bold = True
                p_resumo.add_run(doc_info['resumo'])
                
                # Espaço entre documentos do mesmo procedimento
                doc.add_paragraph("") 
                
            # Separador de procedimentos (fora do loop de documentos)
            doc.add_paragraph("________________________________________________")

        doc.save(os.path.join(pasta_relatorios, "relatorio_pj_docs.docx"))
        return True
    except Exception as e:
        print(f"Erro ao gerar docx: {e}")
        return False
    
# Criação do banco de dados e tabelas
def criar_banco():
    caminho_pasta = "data"
    nome_banco = os.path.join(caminho_pasta, "pj_docs.db")
    
    try:
        if not os.path.exists(caminho_pasta):
            os.makedirs(caminho_pasta)
            print(f"Diretório '{caminho_pasta}' criado.")

        sql_script = """
        PRAGMA foreign_keys = off;
        BEGIN TRANSACTION;

        CREATE TABLE IF NOT EXISTS procedimentos (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT, 
            numero TEXT NOT NULL UNIQUE, 
            descricao TEXT
        );

        CREATE TABLE IF NOT EXISTS documentos (
            id INTEGER NOT NULL PRIMARY KEY AUTOINCREMENT, 
            titulo TEXT NOT NULL, 
            resumo TEXT, 
            data_criacao TEXT DEFAULT CURRENT_TIMESTAMP, 
            procedimento TEXT REFERENCES procedimentos (numero) ON DELETE SET NULL
        );

        COMMIT TRANSACTION;
        PRAGMA foreign_keys = on;
        """
        
        conn = sqlite3.connect(nome_banco)
        cursor = conn.cursor()
        cursor.executescript(sql_script)
        conn.commit()
        
        print(f"Banco de dados '{nome_banco}' operacional!")
        
    except sqlite3.Error as e:
        print(f"Erro ao criar o banco de dados: {e}")
        
    finally:
        if 'conn' in locals() and conn:
            conn.close()

if __name__ == "__main__":
    criar_banco()